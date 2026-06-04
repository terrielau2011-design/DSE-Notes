#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DSE筆記系統 — 匯出筆記工具
用法：python 匯出筆記.py
會讀取同資料夾的「DSE筆記系統.xlsx」，匯出為 Word (.docx) 或 PDF (.pdf)
"""

import os
import sys
from datetime import datetime

try:
    from openpyxl import load_workbook
except ImportError:
    print("❌ 缺少 openpyxl，請先安裝：pip install openpyxl")
    sys.exit(1)

EXCEL_FILE = "DSE筆記系統.xlsx"

# ============================================================
# 設定區（可直接修改）
# ============================================================
EXPORT_SUBJECT = "全部"        # 可選：中史 / 世史 / C&SD / 全部
EXPORT_FORMAT = "兩者皆匯出"    # 可選：Word (.docx) / PDF (.pdf) / 兩者皆匯出
EXPORT_FILTER = "全部"          # 可選：全部 / 只匯出DSE曾出現 / 只匯出熟練度1-2 / 只匯出熟練度1-3
# ============================================================

NOTE_SHEETS = {
    "筆記_中史": "中史",
    "筆記_世史": "世史",
    "筆記_C&SD": "C&SD",
}

PRACTICE_SHEETS = {
    "操練_英文": "英文",
    "操練_數學": "數學",
    "操練_中文": "中文",
}

DSE_TYPE_LABELS = {
    "MC": "🔘 MC選擇題",
    "短答": "✏️ 短答題",
    "DBQ": "📄 DBQ資料題",
    "論述": "📝 論述題",
}

PROFICIENCY_LABELS = {
    1: "⬛ 完全不懂",
    2: "🟥 勉強理解",
    3: "🟨 基本掌握",
    4: "🟩 熟練運用",
    5: "🟦 完全掌握",
}


def read_notes(wb, sheet_name, subject_name, export_filter):
    ws = wb[sheet_name]
    notes = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0]:
            continue
        note = {
            "topic": row[0] or "",
            "subtopic": row[1] or "",
            "content": row[2] or "",
            "keywords": row[3] or "",
            "dse_appeared": str(row[4]) if row[4] else "",
            "dse_types": row[5] or "",
            "dse_year": str(row[6]) if row[6] else "",
            "self_test_q": row[7] or "",
            "self_test_a": row[8] or "",
            "proficiency": row[9] if row[9] else 0,
            "remarks": row[10] if len(row) > 10 and row[10] else "",
        }
        # Apply filter
        if export_filter == "只匯出DSE曾出現" and note["dse_appeared"] != "是":
            continue
        elif export_filter.startswith("只匯出熟練度"):
            try:
                max_level = int(export_filter[-1])
                if note["proficiency"] and int(note["proficiency"]) > max_level:
                    continue
            except (ValueError, IndexError):
                pass
        notes.append(note)
    return notes


def read_practice(wb, sheet_name, subject_name):
    ws = wb[sheet_name]
    records = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0]:
            continue
        record = {
            "date": str(row[0]) if row[0] else "",
            "type": row[1] or "",
            "paper": row[2] or "",
            "source": row[3] or "",
            "dse_year": str(row[4]) if row[4] else "",
            "q_num": row[5] or "",
            "score": row[6] or "",
            "error_type": row[7] or "",
            "error_detail": row[8] or "",
            "improvement": row[9] or "",
            "reviewed": str(row[10]) if len(row) > 10 and row[10] else "",
        }
        records.append(record)
    return records


def format_dse_types(dse_types_str):
    if not dse_types_str:
        return ""
    types = [t.strip() for t in str(dse_types_str).split("/")]
    labels = []
    for t in types:
        if t in DSE_TYPE_LABELS:
            labels.append(DSE_TYPE_LABELS[t])
        elif t:
            labels.append(t)
    return "  ".join(labels)


def format_proficiency(level):
    try:
        return PROFICIENCY_LABELS.get(int(level), str(level))
    except (ValueError, TypeError):
        return str(level)


# ============================================================
# Word (.docx) export
# ============================================================
def export_word(all_data, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("❌ 缺少 python-docx，請先安裝：pip install python-docx")
        return False

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("DSE 筆記匯出", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Export info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"匯出日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}    科目：{EXPORT_SUBJECT}    篩選：{EXPORT_FILTER}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # ---- Note subjects ----
    for sheet_name, subject_name in NOTE_SHEETS.items():
        if EXPORT_SUBJECT != "全部" and subject_name != EXPORT_SUBJECT:
            continue
        notes = all_data.get(f"note_{subject_name}", [])
        if not notes:
            continue

        doc.add_heading(f"📖 {subject_name} 筆記", level=1)

        for i, note in enumerate(notes, 1):
            # Topic heading
            doc.add_heading(f"{i}. {note['topic']} — {note['subtopic']}", level=2)

            # Content
            if note["content"]:
                p = doc.add_paragraph()
                run = p.add_run("📝 筆記重點")
                run.bold = True
                run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
                doc.add_paragraph(note["content"])

            # Keywords
            if note["keywords"]:
                p = doc.add_paragraph()
                run = p.add_run("🔑 關鍵詞：")
                run.bold = True
                run = p.add_run(note["keywords"])
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # DSE info
            if note["dse_appeared"] == "是":
                p = doc.add_paragraph()
                run = p.add_run("📊 DSE出題：")
                run.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                dse_info = f"✅ 已出現"
                if note["dse_types"]:
                    dse_info += f"  |  題型：{format_dse_types(note['dse_types'])}"
                if note["dse_year"]:
                    dse_info += f"  |  年份：{note['dse_year']}"
                run = p.add_run(dse_info)
            elif note["dse_appeared"] == "否":
                p = doc.add_paragraph()
                run = p.add_run("📊 DSE出題：")
                run.bold = True
                run = p.add_run("❌ 未出現")

            # Self-test
            if note["self_test_q"]:
                p = doc.add_paragraph()
                run = p.add_run("❓ 自測題目：")
                run.bold = True
                run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                run = p.add_run(note["self_test_q"])
            if note["self_test_a"]:
                p = doc.add_paragraph()
                run = p.add_run("💡 自測答案：")
                run.bold = True
                run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
                run = p.add_run(note["self_test_a"])

            # Proficiency
            if note["proficiency"]:
                p = doc.add_paragraph()
                run = p.add_run("📊 熟練度：")
                run.bold = True
                run = p.add_run(format_proficiency(note["proficiency"]))

            # Remarks
            if note["remarks"]:
                p = doc.add_paragraph()
                run = p.add_run("📌 備註：")
                run.bold = True
                run = p.add_run(note["remarks"])

            doc.add_paragraph()  # spacer between notes

    # ---- Practice subjects ----
    for sheet_name, subject_name in PRACTICE_SHEETS.items():
        if EXPORT_SUBJECT != "全部" and subject_name != EXPORT_SUBJECT:
            continue
        records = all_data.get(f"practice_{subject_name}", [])
        if not records:
            continue

        doc.add_heading(f"✏️ {subject_name} 操練記錄", level=1)

        # Summary table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["日期", "操練類型", "題號", "得分", "錯誤類型", "已覆檢"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for rec in records:
            row_cells = table.add_row().cells
            row_cells[0].text = rec["date"]
            row_cells[1].text = rec["type"]
            row_cells[2].text = rec["q_num"]
            row_cells[3].text = rec["score"]
            row_cells[4].text = rec["error_type"]
            row_cells[5].text = rec["reviewed"]

        doc.add_paragraph()

        # Error analysis
        errors = [r for r in records if r["error_type"]]
        if errors:
            doc.add_heading("🔍 錯誤分析摘要", level=3)
            for err in errors:
                p = doc.add_paragraph()
                run = p.add_run(f"❌ {err['error_type']}：")
                run.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                run = p.add_run(err["error_detail"])
                if err["improvement"]:
                    p = doc.add_paragraph()
                    run = p.add_run(f"   ✅ 改善行動：")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
                    run = p.add_run(err["improvement"])

        doc.add_paragraph()

    doc.save(output_path)
    print(f"✅ Word 檔案已匯出：{output_path}")
    return True


# ============================================================
# PDF export
# ============================================================
def export_pdf(all_data, output_path):
    try:
        from fpdf import FPDF
    except ImportError:
        print("❌ 缺少 fpdf2，請先安裝：pip install fpdf2")
        return False

    # Find a font that supports Chinese
    font_path = None
    font_bold_path = None
    font_candidates = [
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
         "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
         "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc"),
    ]
    for regular, bold in font_candidates:
        if os.path.exists(regular):
            font_path = regular
            if os.path.exists(bold):
                font_bold_path = bold
            break

    if not font_path:
        import subprocess
        result = subprocess.run(["fc-list", ":lang=zh", "file"], capture_output=True, text=True)
        lines = result.stdout.strip().split("\n")
        for line in lines:
            path = line.split(":")[0].strip()
            if path and os.path.exists(path):
                font_path = path
                break

    if not font_path:
        print("❌ 找不到中文字體，無法匯出 PDF。請安裝中文字體。")
        print("   Ubuntu: sudo apt install fonts-noto-cjk")
        return False

    class PDF(FPDF):
        def header(self):
            self.set_font("chinese", "B", 10)
            self.set_text_color(27, 58, 92)
            self.cell(0, 8, "DSE 筆記匯出", new_x="LMARGIN", new_y="NEXT", align="C")
            self.ln(2)

        def footer(self):
            self.set_y(-15)
            self.set_font("chinese", "", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"第 {self.page_no()}/{{nb}} 頁", align="C")

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.add_font("chinese", "", font_path)
    if font_bold_path:
        pdf.add_font("chinese", "B", font_bold_path)
    else:
        pdf.add_font("chinese", "B", font_path)
    use_chinese = True

    def set_font(size=10, bold=False):
        pdf.set_font("chinese", "B" if bold else "", size)

    pdf.add_page()

    # Title
    set_font(18, True)
    pdf.set_text_color(27, 58, 92)
    pdf.cell(0, 15, "DSE 筆記匯出", new_x="LMARGIN", new_y="NEXT", align="C")

    set_font(8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 6, f"匯出日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  科目：{EXPORT_SUBJECT}  |  篩選：{EXPORT_FILTER}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    # ---- Note subjects ----
    for sheet_name, subject_name in NOTE_SHEETS.items():
        if EXPORT_SUBJECT != "全部" and subject_name != EXPORT_SUBJECT:
            continue
        notes = all_data.get(f"note_{subject_name}", [])
        if not notes:
            continue

        set_font(14, True)
        pdf.set_text_color(139, 69, 19)
        pdf.cell(0, 10, f"[{subject_name}] 筆記", new_x="LMARGIN", new_y="NEXT", align="L")
        pdf.ln(3)

        for i, note in enumerate(notes, 1):
            # Check if we need a new page
            if pdf.get_y() > 250:
                pdf.add_page()

            # Topic
            set_font(12, True)
            pdf.set_text_color(27, 58, 92)
            pdf.cell(0, 8, f"{i}. {note['topic']} — {note['subtopic']}", new_x="LMARGIN", new_y="NEXT", align="L")

            # Content
            if note["content"]:
                set_font(9, True)
                pdf.set_text_color(139, 69, 19)
                pdf.cell(0, 6, "筆記重點：", new_x="LMARGIN", new_y="NEXT")
                set_font(9)
                pdf.set_text_color(50, 50, 50)
                pdf.multi_cell(0, 5, note["content"])

            # Keywords
            if note["keywords"]:
                set_font(9, True)
                pdf.set_text_color(100, 100, 100)
                pdf.cell(0, 6, f"關鍵詞：{note['keywords']}", new_x="LMARGIN", new_y="NEXT")

            # DSE info
            if note["dse_appeared"] == "是":
                set_font(9, True)
                pdf.set_text_color(192, 57, 43)
                dse_info = "DSE出題：已出現"
                if note["dse_types"]:
                    dse_info += f"  |  題型：{note['dse_types']}"
                if note["dse_year"]:
                    dse_info += f"  |  年份：{note['dse_year']}"
                pdf.cell(0, 6, dse_info, new_x="LMARGIN", new_y="NEXT")

            # Self-test
            if note["self_test_q"]:
                set_font(9, True)
                pdf.set_text_color(41, 128, 185)
                pdf.cell(0, 6, f"自測：{note['self_test_q']}", new_x="LMARGIN", new_y="NEXT")
            if note["self_test_a"]:
                set_font(9)
                pdf.set_text_color(39, 174, 96)
                pdf.cell(0, 6, f"答案：{note['self_test_a']}", new_x="LMARGIN", new_y="NEXT")

            # Proficiency
            if note["proficiency"]:
                set_font(9)
                pdf.set_text_color(100, 100, 100)
                pdf.cell(0, 6, f"熟練度：{note['proficiency']}/5", new_x="LMARGIN", new_y="NEXT")

            pdf.ln(4)

    # ---- Practice subjects ----
    for sheet_name, subject_name in PRACTICE_SHEETS.items():
        if EXPORT_SUBJECT != "全部" and subject_name != EXPORT_SUBJECT:
            continue
        records = all_data.get(f"practice_{subject_name}", [])
        if not records:
            continue

        if pdf.get_y() > 230:
            pdf.add_page()

        set_font(14, True)
        pdf.set_text_color(46, 80, 144)
        pdf.cell(0, 10, f"[{subject_name}] 操練記錄", new_x="LMARGIN", new_y="NEXT", align="L")
        pdf.ln(3)

        # Table header
        col_widths = [25, 30, 25, 20, 30, 20, 40]
        headers = ["日期", "類型", "題號", "得分", "錯誤類型", "覆檢", "改善行動"]

        set_font(8, True)
        pdf.set_fill_color(27, 58, 92)
        pdf.set_text_color(255, 255, 255)
        for j, h in enumerate(headers):
            pdf.cell(col_widths[j], 7, h, border=1, align="C", fill=True)
        pdf.ln()

        set_font(8)
        pdf.set_text_color(50, 50, 50)
        for rec in records:
            if pdf.get_y() > 270:
                pdf.add_page()
                # Re-draw header
                set_font(8, True)
                pdf.set_fill_color(27, 58, 92)
                pdf.set_text_color(255, 255, 255)
                for j, h in enumerate(headers):
                    pdf.cell(col_widths[j], 7, h, border=1, align="C", fill=True)
                pdf.ln()
                set_font(8)
                pdf.set_text_color(50, 50, 50)

            row_data = [
                rec["date"][:10],
                str(rec["type"])[:12],
                str(rec["q_num"])[:10],
                str(rec["score"])[:8],
                str(rec["error_type"])[:12],
                rec["reviewed"],
                str(rec["improvement"])[:16],
            ]
            for j, val in enumerate(row_data):
                pdf.cell(col_widths[j], 6, val, new_x="RIGHT", new_y="TOP", align="C")
            pdf.ln()

        pdf.ln(5)

    pdf.output(output_path)
    print(f"✅ PDF 檔案已匯出：{output_path}")
    return True


# ============================================================
# Main
# ============================================================
def main():
    # Read settings from Excel if available
    global EXPORT_SUBJECT, EXPORT_FORMAT, EXPORT_FILTER

    script_dir = os.path.dirname(os.path.abspath(__file__))
    excel_path = os.path.join(script_dir, EXCEL_FILE)

    if not os.path.exists(excel_path):
        print(f"❌ 找不到 {EXCEL_FILE}，請確保檔案在同一資料夾")
        sys.exit(1)

    print(f"📂 讀取檔案：{excel_path}")
    wb = load_workbook(excel_path, data_only=True)

    # Try to read settings from 匯出筆記 sheet
    if "匯出筆記" in wb.sheetnames:
        ws_export = wb["匯出筆記"]
        val_b4 = ws_export["B4"].value
        val_b5 = ws_export["B5"].value
        val_b6 = ws_export["B6"].value
        valid_subjects = ["中史", "世史", "C&SD", "全部"]
        valid_formats = ["Word (.docx)", "PDF (.pdf)", "兩者皆匯出"]
        valid_filters = ["全部", "只匯出DSE曾出現", "只匯出熟練度1-2", "只匯出熟練度1-3"]
        if val_b4 and val_b4 in valid_subjects:
            EXPORT_SUBJECT = val_b4
        if val_b5 and val_b5 in valid_formats:
            EXPORT_FORMAT = val_b5
        if val_b6 and val_b6 in valid_filters:
            EXPORT_FILTER = val_b6

    print(f"⚙️ 匯出設定：科目={EXPORT_SUBJECT}  格式={EXPORT_FORMAT}  篩選={EXPORT_FILTER}")
    print()

    # Read all data
    all_data = {}
    for sheet_name, subject_name in NOTE_SHEETS.items():
        if EXPORT_SUBJECT != "全部" and subject_name != EXPORT_SUBJECT:
            continue
        all_data[f"note_{subject_name}"] = read_notes(wb, sheet_name, subject_name, EXPORT_FILTER)
        count = len(all_data[f"note_{subject_name}"])
        print(f"  📖 {subject_name}：{count} 條筆記")

    for sheet_name, subject_name in PRACTICE_SHEETS.items():
        if EXPORT_SUBJECT != "全部" and subject_name != EXPORT_SUBJECT:
            continue
        all_data[f"practice_{subject_name}"] = read_practice(wb, sheet_name, subject_name)
        count = len(all_data[f"practice_{subject_name}"])
        print(f"  ✏️ {subject_name}：{count} 條操練記錄")

    if not any(all_data.values()):
        print("❌ 沒有找到任何筆記或操練記錄")
        sys.exit(1)

    print()

    # Export
    date_str = datetime.now().strftime("%Y%m%d_%H%M")

    if EXPORT_FORMAT in ["Word (.docx)", "兩者皆匯出"]:
        docx_name = f"DSE筆記_{EXPORT_SUBJECT}_{date_str}.docx"
        docx_path = os.path.join(script_dir, docx_name)
        export_word(all_data, docx_path)

    if EXPORT_FORMAT in ["PDF (.pdf)", "兩者皆匯出"]:
        pdf_name = f"DSE筆記_{EXPORT_SUBJECT}_{date_str}.pdf"
        pdf_path = os.path.join(script_dir, pdf_name)
        export_pdf(all_data, pdf_path)

    print()
    print("🎉 匯出完成！")


if __name__ == "__main__":
    main()
